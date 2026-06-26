"""
Konvertuje UPUTSTVO_ZARADE_KOMPLETNO.md u lepo formatiran Word dokument.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

# ── Boje (hex stringovi za XML, RGBColor za font) ─────────────────────────────
HEX_DARK_GREEN  = "1B5E20"
HEX_MED_GREEN   = "2E7D32"
HEX_LIGHT_GREEN = "43A047"
HEX_WHITE       = "FFFFFF"
HEX_GRAY_BG     = "E8F5E9"
HEX_ALT_ROW     = "E8F5E9"
HEX_NORM_ROW    = "FFFFFF"
HEX_CODE_BG     = "F5F5F5"
HEX_WARN_BG     = "FFF3E0"
HEX_INFO_BG     = "E3F2FD"

C_DARK_GREEN    = RGBColor(0x1B, 0x5E, 0x20)
C_MED_GREEN     = RGBColor(0x2E, 0x7D, 0x32)
C_LIGHT_GREEN   = RGBColor(0x43, 0xA0, 0x47)
C_BLUE          = RGBColor(0x15, 0x65, 0xC0)
C_ORANGE        = RGBColor(0xE6, 0x51, 0x00)
C_WHITE         = RGBColor(0xFF, 0xFF, 0xFF)
C_GRAY_TEXT     = RGBColor(0x37, 0x47, 0x4F)
C_CODE_TEXT     = RGBColor(0x1A, 0x23, 0x7E)
C_INFO_TEXT     = RGBColor(0x0D, 0x47, 0xA1)

# ── XML pomocne funkcije ───────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, color="B0BEC5"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def add_para_border_bottom(para, color="C8E6C9"):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def set_para_shading(para, hex_color: str):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)

# ── Inline tekst (bold, code) ─────────────────────────────────────────────────

def add_inline_text(para, text, base_size=10.5, base_color=C_GRAY_TEXT):
    segments = re.split(r'`([^`]+)`', text)
    for i, seg in enumerate(segments):
        if i % 2 == 1:
            run = para.add_run(seg)
            run.font.name = "Courier New"
            run.font.size = Pt(9.5)
            run.font.color.rgb = C_BLUE
            run.bold = True
        else:
            parts = re.split(r'\*\*(.+?)\*\*', seg)
            for j, part in enumerate(parts):
                if not part:
                    continue
                run = para.add_run(part)
                run.font.name = "Calibri"
                run.font.size = Pt(base_size)
                if j % 2 == 1:
                    run.bold = True
                    run.font.color.rgb = C_DARK_GREEN
                else:
                    run.font.color.rgb = base_color

# ── Margine i stilovi ─────────────────────────────────────────────────────────

def setup_doc(doc: Document):
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.0)

# ── Naslov dokumenta ──────────────────────────────────────────────────────────

def add_title_block(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    set_para_shading(p, HEX_GRAY_BG)
    add_para_border_bottom(p, HEX_MED_GREEN)
    run = p.add_run("UPUTSTVO ZA RAD SA ZARADAMA")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = C_DARK_GREEN

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(2)
    set_para_shading(p2, HEX_GRAY_BG)
    r2 = p2.add_run("Algoritam — Program za obračun zarada i ličnih primanja")
    r2.font.name = "Calibri"; r2.font.size = Pt(13)
    r2.font.color.rgb = C_MED_GREEN; r2.italic = True

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(12)
    set_para_shading(p3, HEX_GRAY_BG)
    r3 = p3.add_run("Verzija: jun 2026  |  Sa konkretnim primerima podataka za unos")
    r3.font.name = "Calibri"; r3.font.size = Pt(10)
    r3.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

# ── Heading funkcije ──────────────────────────────────────────────────────────

def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after  = Pt(4)
    set_para_shading(p, HEX_DARK_GREEN)
    run = p.add_run("  " + text)
    run.font.name = "Calibri"; run.font.size = Pt(14)
    run.bold = True; run.font.color.rgb = C_WHITE

def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(3)
    set_para_shading(p, "C8E6C9")
    add_para_border_bottom(p, HEX_MED_GREEN)
    run = p.add_run(text)
    run.font.name = "Calibri"; run.font.size = Pt(12)
    run.bold = True; run.font.color.rgb = C_MED_GREEN

def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.name = "Calibri"; run.font.size = Pt(11)
    run.bold = True; run.font.color.rgb = C_LIGHT_GREEN
    add_para_border_bottom(p, "A5D6A7")

def add_h4(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run(text)
    run.font.name = "Calibri"; run.font.size = Pt(10.5)
    run.bold = True; run.font.color.rgb = C_BLUE

# ── Tabele ────────────────────────────────────────────────────────────────────

def parse_md_table(lines, start):
    rows = []
    i = start
    while i < len(lines):
        line = lines[i].strip()
        if not line.startswith('|'):
            break
        if re.match(r'^\|[-| :]+\|$', line):
            i += 1; continue
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)
        i += 1
    return rows, i

def add_md_table(doc, rows):
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    for ri, row_data in enumerate(rows):
        is_hdr = (ri == 0)
        tr = tbl.rows[ri]
        for ci in range(ncols):
            cell = tr.cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            text = row_data[ci] if ci < len(row_data) else ""
            if is_hdr:
                set_cell_bg(cell, HEX_DARK_GREEN)
            elif ri % 2 == 0:
                set_cell_bg(cell, HEX_ALT_ROW)
            else:
                set_cell_bg(cell, HEX_NORM_ROW)
            set_cell_border(cell)
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after  = Pt(2)
            if is_hdr:
                run = para.add_run(text)
                run.font.name = "Calibri"; run.font.size = Pt(9.5)
                run.bold = True; run.font.color.rgb = C_WHITE
            else:
                add_inline_text(para, text, base_size=9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── Code blok ─────────────────────────────────────────────────────────────────

def add_code_block(doc, code_lines):
    for line in code_lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Cm(0.5)
        set_para_shading(p, HEX_CODE_BG)
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"; run.font.size = Pt(9)
        run.font.color.rgb = C_CODE_TEXT
    # mala razmak posle bloka
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after  = Pt(6)

# ── Bullet lista ──────────────────────────────────────────────────────────────

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    add_inline_text(p, text)

# ── Upozorenje/napomena blok ──────────────────────────────────────────────────

def add_warning_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    set_para_shading(p, HEX_WARN_BG)
    add_para_border_bottom(p, "E65100")
    run = p.add_run("  " + text)
    run.font.name = "Calibri"; run.font.size = Pt(10.5)
    run.font.color.rgb = C_ORANGE; run.bold = True

def add_info_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    set_para_shading(p, HEX_INFO_BG)
    add_inline_text(p, text, base_size=10, base_color=C_INFO_TEXT)

# ── Glavni parser ─────────────────────────────────────────────────────────────

def build_doc(doc: Document, md_path: str):
    with open(md_path, encoding='utf-8') as f:
        lines = [l.rstrip('\n') for l in f]

    i = 0
    skip_toc = False  # preskoči sadržaj (numerisane linkove)

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # ── Naslov dokumenta (H1 sa meta podacima) ──
        if line.startswith('# ') and not line.startswith('## '):
            text = line[2:].strip()
            if any(x in text for x in ['UPUTSTVO', 'Algoritam', 'Verzija']):
                i += 1; continue
            add_h1(doc, text)
            i += 1; continue

        # ── H2 (## ) ──
        if line.startswith('## ') and not line.startswith('### '):
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', line[3:].strip())
            if re.match(r'^\d+\.', text):
                add_h1(doc, text)
            else:
                add_h2(doc, text)
            i += 1; continue

        # ── H3 (### ) ──
        if line.startswith('### ') and not line.startswith('#### '):
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', line[4:].strip())
            add_h2(doc, text)
            i += 1; continue

        # ── H4 (#### ) ──
        if line.startswith('#### '):
            text = line[5:].strip()
            add_h3(doc, text)
            i += 1; continue

        # ── Code blok ──
        if stripped.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1
            add_code_block(doc, code_lines)
            continue

        # ── Tabela ──
        if stripped.startswith('|'):
            rows, i = parse_md_table(lines, i)
            add_md_table(doc, rows)
            continue

        # ── Bullet lista ──
        m = re.match(r'^(\s*)[-*] (.+)', line)
        if m:
            lvl = len(m.group(1)) // 2
            add_bullet(doc, m.group(2), level=lvl)
            i += 1; continue

        m2 = re.match(r'^(\s*)\d+\. (.+)', line)
        if m2:
            # Sadrzaj linkovi — preskoci
            if re.match(r'^\d+\. \[', stripped):
                i += 1; continue
            lvl = len(m2.group(1)) // 2
            add_bullet(doc, m2.group(2), level=lvl)
            i += 1; continue

        # ── Horizontalna linija ──
        if re.match(r'^---+$', stripped):
            p = doc.add_paragraph()
            add_para_border_bottom(p, HEX_MED_GREEN)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(8)
            i += 1; continue

        # ── Blockquote ──
        if stripped.startswith('> '):
            inner = stripped[2:]
            # VAŽNO / Napomena = narandžasto
            if '**VAŽNO' in inner or '**Napomena' in inner or 'NIKADA' in inner:
                clean = re.sub(r'\*\*(.+?)\*\*', r'\1', inner)
                add_warning_block(doc, clean)
            else:
                clean = re.sub(r'\*\*(.+?)\*\*', r'\1', inner)
                add_info_block(doc, clean)
            i += 1; continue

        # ── Prazna linija ──
        if not stripped:
            i += 1; continue

        # ── Obican paragraf ──
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(3)
        add_inline_text(p, stripped)
        i += 1

# ── Entry point ───────────────────────────────────────────────────────────────

def main():
    input_path  = r"c:\Workspace\algoritam-migration\newproject\zarade_uputstva\UPUTSTVO_ZARADE_KOMPLETNO.md"
    output_path = r"c:\Workspace\algoritam-migration\newproject\zarade_uputstva\UPUTSTVO_ZARADE_KOMPLETNO.docx"

    doc = Document()
    setup_doc(doc)
    add_title_block(doc)
    build_doc(doc, input_path)
    doc.save(output_path)
    print(f"OK: {output_path}")

if __name__ == '__main__':
    main()
